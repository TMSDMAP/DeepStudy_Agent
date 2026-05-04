# -*- coding: utf-8 -*-
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

TEMPLATE = Path(os.environ["REPORT_TEMPLATE"])
OUT = Path(os.environ["REPORT_OUT"])

WORK_TITLE = "深学讲义Agent：面向课堂教学的多智能体讲义生成与智能评价系统"
DATE_TEXT = "2026年4月19日"


def set_run_font(run, size=11, bold=False, color=None, name="宋体"):
    run.font.name = name
    run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_para(doc, text="", size=11, bold=False, align=None, first_line=True, spacing_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = 1.25
    pf.space_after = Pt(spacing_after)
    if first_line and text:
        pf.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_heading(doc, text, level=1):
    style = f"Heading {level}"
    try:
        p = doc.add_paragraph(style=style)
    except Exception:
        p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=16 if level == 1 else 14 if level == 2 else 12, bold=True, name="黑体")
    return p


def add_table(doc, headers, rows, caption=None):
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        set_run_font(run, size=10, bold=True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for para in hdr[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                set_run_font(run, size=10, bold=True, name="黑体")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cells[i].paragraphs:
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    set_run_font(run, size=10)
    doc.add_paragraph()
    return table


def clear_body(doc):
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith('sectPr'):
            continue
        body.remove(child)


def add_cover(doc):
    for _ in range(2):
        add_para(doc, "", first_line=False)
    p = add_para(doc, "2026年（第19届）", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, spacing_after=4)
    p = add_para(doc, "中国大学生计算机设计大赛", size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, spacing_after=18)
    p = add_para(doc, "人工智能实践赛作品报告", size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, spacing_after=36)
    add_para(doc, "作品编号：待填写", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, spacing_after=12)
    add_para(doc, f"作品名称：{WORK_TITLE}", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, spacing_after=12)
    add_para(doc, f"填写日期：{DATE_TEXT}", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, spacing_after=24)
    add_para(doc, "说明：本报告为参赛作品草稿版，作品编号、团队成员、最终测试数据等信息需在提交前根据赛事平台实际信息补充。", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    doc.add_page_break()


def add_contents(doc):
    add_heading(doc, "目  录", 1)
    items = [
        "第1章 作品概述",
        "第2章 问题分析",
        "  2.1 问题来源",
        "  2.2 现有解决方案",
        "  2.3 本作品要解决的痛点问题",
        "  2.4 解决问题的思路",
        "第3章 技术方案",
        "第4章 系统实现",
        "第5章 测试分析",
        "第6章 作品总结",
        "  6.1 作品特色与创新点",
        "  6.2 应用推广",
        "  6.3 作品展望",
        "参考文献",
    ]
    for item in items:
        add_para(doc, item, first_line=False, spacing_after=2)
    add_para(doc, "注：正式提交前可在 Word 中更新为自动目录。", size=10, first_line=False)
    doc.add_page_break()


def build_report(doc):
    add_heading(doc, "第1章 作品概述", 1)
    add_para(doc, f"{WORK_TITLE}是一套面向教师备课、课堂讲授和课后评价的人工智能教学辅助系统。作品的创意来源于实际教学准备中的高频问题：教师往往拥有课件、教材、论文、网页、题库和学生作答等多种材料，但通用大模型通常只能给出问答式回复，难以稳定产出符合课堂结构、排版规范和教学闭环要求的讲义、PPT、网页演示与习题解析。")
    add_para(doc, "本作品以“从问答助手升级为课堂内容生产流水线”为核心目标，引入多Agent协同机制，将讲义生成过程拆解为主题教研、教学设计、互动活动、题库生成、总编整合、质量检查和代码运行验证等多个角色。系统支持教师上传本次临时资料，也支持将资料长期沉淀到个人知识库，在生成时自动检索并注入相关片段，使讲义更贴合教师已有材料，而不是完全依赖模型的通用知识。")
    add_para(doc, "作品的主要用户群体包括高校教师、中小学信息技术教师、助教、教研团队以及需要制作课程材料的学生。系统当前重点服务于“给定主题和教师资料，生成可直接用于课堂的结构化讲义”这一场景，同时保留智能作答与批改引擎，用于上传题目或学生作答图片后进行题干数字化、分步批改、变式生成和知识库沉淀。")
    add_para(doc, "系统的主要功能包括：多Agent讲义生成、临时外部资料解析、个人长期知识库、PDF/PPT/DOCX/TXT/HTML/图片等多格式输入、Markdown/Word/PDF/PPT/HTML多格式导出、选择题与答案解析规范化、示例代码运行检查、PPT主题切换、知识库检索自测以及长期知识库清空等。作品的应用价值在于降低教师备课成本，提高讲义结构质量，并将一次性生成内容转化为可复用的教学知识资产。")

    add_heading(doc, "第2章 问题分析", 1)
    add_heading(doc, "2.1 问题来源", 2)
    add_para(doc, "随着生成式人工智能进入教学场景，教师可以借助大模型快速生成解释、题目和课堂活动。然而在实际使用中，通用问答工具往往缺少课程意识、资料意识和格式意识：它能回答“词法分析是什么”，却不一定能根据教师上传的PPT生成一节45分钟课堂讲义；它能列出题目，却不一定能保证选项换行、答案解析一一对应；它能写代码，却不一定会检查代码是否完整可运行。")
    add_para(doc, "本作品的问题来源正是这些实际落差：教师需要的是完整教学交付物，而不是零散回答；需要模型遵循课程材料，而不是随意扩展；需要系统能导出PPT和HTML并适配课堂展示，而不是只生成一段文本。")

    add_heading(doc, "2.2 现有解决方案", 2)
    add_para(doc, "现有方案大致可以分为四类：通用大模型问答工具、在线课件生成工具、题库/作业批改系统和知识库问答系统。它们各有优势，但在“基于教师资料生成完整课堂讲义”这一综合任务上仍存在不足。")
    add_table(doc, ["方案类型", "代表能力", "主要不足"], [
        ["通用大模型问答", "解释概念、生成文本、编写简单代码", "缺少稳定教学结构，容易偏离教师上传资料，输出格式不可控"],
        ["课件生成工具", "快速生成PPT页面和视觉模板", "重展示轻教学，缺少课堂流程、互动设计和质量检查"],
        ["题库与批改系统", "自动出题、客观题批改、部分主观题评价", "与讲义生成和教师知识库割裂，难以形成完整教学闭环"],
        ["知识库问答系统", "基于文档检索回答问题", "偏向检索问答，缺少多Agent协作、导出和教学设计能力"],
    ], "表 1  现有解决方案对比")

    add_heading(doc, "2.3 本作品要解决的痛点问题", 2)
    add_para(doc, "第一，通用AI输出缺少课堂结构。教学讲义不仅需要知识点，还需要学习目标、学情分析、课堂流程、互动活动、习题、答案解析和课后延伸等完整环节。")
    add_para(doc, "第二，教师资料难以被有效利用。教师往往已经有PPT、PDF论文、教材摘录和个人讲义，如果系统不能解析并检索这些资料，生成内容就容易变成泛泛而谈。")
    add_para(doc, "第三，生成结果缺少质量约束。实际项目中曾出现选择题答案多于题目、答案解析残缺、代码块无法识别、PPT页面文字溢出等问题，因此需要在生成后增加本地质量闸门和格式规范化。")
    add_para(doc, "第四，教学产物需要多格式交付。教师真实场景中既要可编辑Word，也要课堂投屏PPT，还可能需要HTML网页演示和PDF归档，单一文本输出无法满足需求。")

    add_heading(doc, "2.4 解决问题的思路", 2)
    add_para(doc, "本作品采用“资料解析 + 知识检索 + 多Agent协同 + 本地质量闸门 + 多格式导出”的整体思路。教师输入教学主题和硬性要求后，系统先解析临时资料并检索长期知识库，再由不同Agent分别完成教研拆解、教学设计、互动活动和题库设计，最后由总编Agent融合为讲义，并由质检Agent和本地规则检查输出完整性。")
    add_table(doc, ["数据类型", "来源", "用途", "样例"], [
        ["PPTX课件", "教师上传", "提取课程主线和关键概念", "词法分析(1).pptx"],
        ["PDF/DOCX/TXT/HTML", "教师资料、教材摘录、论文", "作为外部知识库补充事实依据", "课程讲义、论文、教材片段"],
        ["图片", "题目截图、学生作答", "多模态识别、批改和变式生成", "手写解题过程或试题截图"],
        ["长期知识库切片", "教师多次入库资料", "跨会话召回历史教学材料", "个人知识库中的文本片段"],
    ], "表 2  系统使用的数据来源与用途")

    add_heading(doc, "第3章 技术方案", 1)
    add_para(doc, "系统整体技术路线如图1所示：教师输入主题和资料，资料解析模块将不同格式转化为文本切片，知识库模块进行检索召回，多Agent工作流分别完成教研、教学设计、互动设计、题库设计、内容整合和质量评估，最后由导出模块生成Markdown、Word、PDF、PPT和HTML。")
    p = add_para(doc, "教师输入/上传资料 → 文档解析与切片 → 临时资料与长期知识库检索 → 教研Agent → 教学设计Agent → 互动设计Agent → 题库Agent → 总编Agent → 质检Agent/本地质量闸门 → 多格式导出", size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_para(doc, "图 1  系统技术路线框架图", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_para(doc, "多Agent协同是本作品的核心。教研Agent负责主题拆解、知识检索和大纲构建；教学设计Agent负责学习目标、学情假设、时间安排和分层策略；互动设计Agent负责编排小组讨论、课堂提问和过程性评价；题库Agent负责生成分层选择题和答案解析；总编Agent负责统一语言、结构和章节；质检Agent负责检查主题范围、章节完整性、题库对应关系、代码块完整性和占位符残留。")
    add_para(doc, "知识库模块分为本次临时资料和个人长期知识库。临时资料只在本次生成中注入，适合教师上传某一节课的PPT或论文；长期知识库支持多次追加、检索自测和一键清空，适合沉淀教师自己的课程资料。PDF解析采用pypdf优先、pdfplumber兜底的策略，并在界面显示解析状态、文本长度和失败原因，使教师能够判断资料是否真的被系统识别。")
    add_para(doc, "质量控制采用模型评价与本地规则结合的方式。模型评价负责语义层面的合理性判断，本地规则负责硬性约束，例如讲义是否过短、是否缺少主要章节、是否存在“待补充答案”或TODO、选择题是否少于4道、答案是否少于4条、要求代码时是否存在Python代码块等。若模型最终草稿不完整，系统会使用各子Agent结果进行本地重组，避免长时间等待后只得到残缺内容。")
    add_para(doc, "多模态智能作答与批改引擎使用兼容OpenAI接口的多模态模型服务，可接入DashScope等支持图像输入的模型，用于识别题目截图或学生作答图片。该模块将识别、求解、批改、变式生成和知识库沉淀串联起来，为课后练习和个性化辅导提供基础。")

    add_heading(doc, "第4章 系统实现", 1)
    add_para(doc, "系统采用Python实现，前端交互使用Streamlit，Agent编排使用LangGraph/LangChain思想组织工作流，大模型调用采用OpenAI兼容接口，便于在DeepSeek、DashScope等不同服务之间切换。文档解析和导出模块使用python-pptx、pypdf、pdfplumber、python-docx、ReportLab等组件。")
    add_para(doc, "系统界面分为两个核心引擎：其一是讲义生成引擎，教师输入教学主题、教学风格和硬性要求，上传本次临时资料或检索长期知识库后，系统生成讲义并提供多格式下载；其二是智能作答与批改引擎，教师上传题目或学生作答图片后，系统进行题干识别、思路分析、分步批改和结果沉淀。")
    add_para(doc, "讲义生成流程包括六个关键步骤：第一，读取主题、教师要求、代码和题库开关；第二，解析临时资料并记录解析状态；第三，检索长期知识库并将命中片段注入上下文；第四，执行多Agent工作流；第五，进行Markdown规范化、题目选项换行、占位符清理和质量检查；第六，生成Markdown、Word、PDF、PPT和HTML导出文件。")
    add_para(doc, "在工程迭代过程中，系统针对真实使用问题进行了多次修正。例如，针对PPT导出段落为空导致的IndexError，增加了安全获取文本run的函数；针对题库中出现多余“待补充答案”的问题，增加题目与答案数量一致性检查；针对HTML和PPT文字超出页面的问题，增加内容拆分和分页逻辑；针对PDF资料无法识别的问题，增加pypdf与pdfplumber双解析方案和失败原因提示。")
    add_table(doc, ["模块", "实现内容", "对应用户价值"], [
        ["资料解析", "支持PPTX、PDF、DOCX、TXT、HTML和图片解析", "让教师已有资料成为生成依据"],
        ["知识库", "临时资料注入、长期知识库追加、检索自测、清空", "支持持续沉淀个人课程资产"],
        ["多Agent工作流", "教研、教学设计、互动、题库、总编、质检", "提升讲义结构完整性"],
        ["导出模块", "Markdown、Word、PDF、PPT、HTML", "满足备课、投屏和归档需求"],
        ["批改引擎", "图片识别、智能作答、分步批改、变式生成", "服务课后练习与个性化反馈"],
    ], "表 3  系统模块实现概览")

    add_heading(doc, "第5章 测试分析", 1)
    add_para(doc, "当前测试以功能验证和典型案例验证为主，重点确认系统是否能够完成从资料上传到讲义导出的完整链路。测试环境为本地Windows环境，Python虚拟环境运行Streamlit应用。测试数据包括课程PPT、文学类文本资料、PDF论文和手工输入的教学约束。")
    add_table(doc, ["测试项", "测试方法", "阶段性结果"], [
        ["PPTX临时资料解析", "上传词法分析(1).pptx，检查解析状态", "解析成功，提取文本739字符，可作为本次外部资料注入"],
        ["长期知识库", "上传文本资料后进行关键词检索自测", "可以命中并展示召回片段，支持清空知识库"],
        ["PDF解析", "使用文献PDF测试pypdf和pdfplumber兜底", "依赖安装后可显示解析方法、文本长度和失败原因"],
        ["题库规范", "生成至少4道选择题并检查选项与答案", "增加A/B/C/D换行、答案数量检查和占位符拦截"],
        ["PPT/HTML导出", "对长段落、表格和代码块进行导出测试", "增加分页和文本清理，减少单页溢出"],
        ["代码验证", "检测讲义中的Python代码块并可触发运行检查", "可报告代码块数量和运行状态"],
    ], "表 4  阶段性测试结果")
    add_para(doc, "测试表明，系统已具备完整的课堂讲义生成闭环，能够处理教师上传资料、生成结构化讲义并导出多种格式。同时，测试也暴露出大模型输出不稳定、资料解析依赖环境、知识库召回精度需要提升等问题。针对这些问题，作品已加入本地质量闸门、解析状态提示和兜底重组机制，后续仍需通过更多课程样例和教师试用数据进行量化评估。")
    add_para(doc, "拟在后续决赛阶段补充的量化指标包括：资料解析成功率、讲义章节完整率、题库答案对应准确率、PPT页面溢出率、教师人工修改时长、用户满意度和多轮生成平均耗时。这些指标将用于进一步验证系统在真实教学场景中的有效性和稳定性。")

    add_heading(doc, "第6章 作品总结", 1)
    add_heading(doc, "6.1 作品特色与创新点", 2)
    add_para(doc, "本作品的第一项特色是将通用问答式AI升级为面向课堂交付物的多Agent生产流程。系统不是简单回答教师问题，而是围绕教学目标、课堂流程、互动活动、题库、代码和导出格式形成完整闭环。")
    add_para(doc, "第二项特色是将本次临时资料和个人长期知识库结合。教师既可以针对某节课上传PPT或PDF，也可以长期积累个人课程资料，使系统逐步贴合教师自己的教学风格和知识体系。")
    add_para(doc, "第三项特色是强调质量闸门和工程可用性。系统针对占位符残留、答案解析缺失、代码块识别失败、PPT文字溢出等真实问题设置规则检查，减少大模型不稳定输出对课堂使用的影响。")
    add_para(doc, "第四项特色是保留多模态智能作答与批改能力，使系统不仅服务课前备课，也能延伸到课中互动和课后评价。")

    add_heading(doc, "6.2 应用推广", 2)
    add_para(doc, "本作品可首先推广到高校计算机类课程，例如编译原理、数据结构、人工智能导论、机器学习和软件工程等。这类课程知识结构清晰、代码示例较多、教师资料丰富，适合作为系统试点。随后可拓展到语文、历史、物理等需要文本解读和课堂活动设计的课程。")
    add_para(doc, "推广方式上，可以采用本地部署和云端部署两种模式。本地部署适合重视资料隐私的教师和教研组，云端部署适合学校统一使用。系统的OpenAI兼容接口设计也便于根据成本、速度和多模态能力选择不同模型服务。")

    add_heading(doc, "6.3 作品展望", 2)
    add_para(doc, "未来作品将从四个方向继续完善：第一，提升知识库召回质量，引入更稳定的向量检索、章节定位和重排序机制；第二，增加对扫描版PDF和手写材料的OCR支持；第三，构建更多课程模板和评价量表，使不同学科生成结果更符合教学规范；第四，加入教师反馈闭环，让系统能够根据教师修改记录持续优化生成策略。")
    add_para(doc, "总体来看，深学讲义Agent已经从单纯问答助手发展为具备资料解析、多Agent协同、质量控制、多格式导出和智能批改能力的教学辅助系统。随着测试样本和教师反馈的增加，作品有望成为教师备课、课堂展示和课后评价的一体化AI工具。")

    add_heading(doc, "参考文献", 1)
    refs = [
        "[1] Brown T B, Mann B, Ryder N, et al. Language Models are Few-Shot Learners[C]//Advances in Neural Information Processing Systems, 2020.",
        "[2] Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]//Advances in Neural Information Processing Systems, 2020.",
        "[3] OpenAI. GPT-4 Technical Report[EB/OL]. arXiv:2303.08774, 2023.",
        "[4] LangChain Contributors. LangChain Documentation[EB/OL].",
        "[5] Streamlit Inc. Streamlit Documentation[EB/OL].",
        "[6] Microsoft. Office Open XML File Formats Documentation[EB/OL].",
    ]
    for ref in refs:
        add_para(doc, ref, first_line=False, spacing_after=3)
    add_para(doc, "注：正式提交前建议根据实际采用的模型、框架版本和赛事要求补全访问日期与引用格式。", size=10, first_line=False)


def main():
    doc = Document(str(TEMPLATE))
    clear_body(doc)
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.8)
    try:
        normal = doc.styles['Normal']
        normal.font.name = '宋体'
        normal._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", '宋体')
        normal.font.size = Pt(11)
    except Exception:
        pass
    add_cover(doc)
    add_contents(doc)
    build_report(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(str(OUT))

if __name__ == "__main__":
    main()
